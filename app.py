import os
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.filechooser import FileChooserListView
from kivy.uix.spinner import Spinner
from kivy.uix.popup import Popup
from kivy.uix.textinput import TextInput
from cryptography.fernet import Fernet
from kivy.core.clipboard import Clipboard
from spacy.pipeline import EntityRuler
from docx import Document
import platform
from io import StringIO
import re
import pandas as pd
import os
from abc import ABC, abstractmethod
import fitz
import spacy # type: ignore
import stanza
import sys
import subprocess
import time

# Anonymization Strategy Interface
class AnonymizationStrategy(ABC):
    @abstractmethod
    def modify_text(self, text):
        pass
    def normalize_spaced_text(self, text):
        # Replace multiple spaces with a placeholder
        text = re.sub(r'\s{2,}', '||', text)
        # Reduce any remaining single spaces (that aren't part of multiple spaces)
        text = re.sub(r'(?<=\b\w) (?=\w\b)', '', text)
        # Reintroduce the spaces for word boundaries
        text = re.sub(r'\|\|', ' ', text)
        return text
    def contains_one_letter_space_one_letter(self, text):
        pattern = r'\b\w\s\w\b'
        return re.search(pattern, text) is not None

# Concrete Strategy using spaCy
class SpacyAnonymization(AnonymizationStrategy):
    def __init__(self):
        self.nlp = spacy.load('en_core_web_lg')

    def modify_text(self, text):
        # patterns = [{"label": "IGNORE", "pattern": "QTY"}]
        # # self.nlp.remove_pipe("entity_ruler")
        # if 'entity_ruler' in self.nlp.pipe_names:
        #     self.nlp.remove_pipe('entity_ruler')
        # ruler = self.nlp.add_pipe("entity_ruler")
        # ruler.add_patterns(patterns)
        if self.contains_one_letter_space_one_letter(text):
            text = self.normalize_spaced_text(text)
        doc = self.nlp(text)
        entities = doc.ents
        sorted_entities = sorted(entities, key=lambda e: e.start_char, reverse=True)
        anonymized_text = text


        for entity in sorted_entities:
            # Using entity type directly for replacement
            entity_type = entity.label_
            if entity_type=='CARDINAL':
                entity_type="NUMBER"
            # Read the span from doc.ents and apply anonymization to that part of the text
            anonymized_text = anonymized_text[:entity.start_char] + "{{"+ entity_type + "}}" + anonymized_text[entity.end_char:]
    
        return anonymized_text

# Concrete Strategy using Stanza
class StanzaAnonymization(AnonymizationStrategy):
    def __init__(self):
        self.nlp = stanza.Pipeline(lang='en', processors='tokenize,ner')

    def modify_text(self, text):
        if self.contains_one_letter_space_one_letter(text):
            text = self.normalize_spaced_text(text)
        doc = self.nlp(text)
        entities = doc.ents
        sorted_entities = sorted(entities, key=lambda e: e.start_char, reverse=True)
        anonymized_text = text

        for entity in sorted_entities:
            # Using entity type directly for replacement
            entity_type = entity.type
            if entity_type=='CARDINAL':
                entity_type="NUMBER"
            # Read the span from doc.ents and apply anonymization to that part of the text
            anonymized_text = anonymized_text[:entity.start_char] + "{{"+ entity_type + "}}" + anonymized_text[entity.end_char:]

        return anonymized_text

class DocumentHandler(ABC):
    def __init__(self, anonymizer: AnonymizationStrategy):
        self.anonymizer = anonymizer
        self.key = Fernet.generate_key().decode()

    @abstractmethod
    def process_document(self, input_path, output_path):
        pass

    @abstractmethod
    def deprocess_document(self, dataframe, output_path):
        pass
    
    def uuid_path(self, pdf_path):
        # Split the path into directory and filename
        directory, filename = os.path.split(pdf_path)
        # Split the filename into name and extension
        name, extension = os.path.splitext(filename)
        # Construct the new filename with UUID
        new_filename = f"{name}_uuid{extension}.txt"
        # Construct the new path
        new_path = os.path.join(directory, new_filename)
        return new_path

class PDFHandler(DocumentHandler):
    def convert_color(self, color_int):
        r = (color_int >> 16) & 0xff
        g = (color_int >> 8) & 0xff
        b = color_int & 0xff
        return (r / 255.0, g / 255.0, b / 255.0)  # Convert to [0, 1] range

    def deprocess_document(self, dataframe, output_path):
        start_time = time.time()
        new_doc = fitz.open()
        new_page = new_doc.new_page()
        for index, row in dataframe.iterrows():
            rect = fitz.Rect(row['bbox']) 
            new_page.insert_text(rect.tl, row['text'], fontsize=row['size'], fontname=row['font'], color=row['color'])
        new_doc.save(output_path)
        end_time = time.time()  
        elapsed_time = (end_time - start_time) * 1000  
        print(f"Deanonymization process for pdf took {elapsed_time:.2f} milliseconds")

    def process_document(self, input_path, output_path):
        start_time = time.time()
        doc = fitz.open(input_path)
        text_data = []
        for page in doc:
            text_dict = page.get_text("dict")
            print(text_dict)
            for block in text_dict['blocks']:
                if block['type'] == 0:  # Type 0 is text
                    for line in block['lines']:
                        for span in line['spans']:
                            # Check if the text is bold
                            if 'Bold' in span['font'] or span['flags'] & 2:
                                font_name = "Helvetica-Bold"
                            else:
                                font_name = "Helvetica"
                            span_data = {
                                'anontext': self.anonymizer.modify_text(span['text']),
                                'text': span['text'],
                                'bbox': span['bbox'],
                                'size': span['size'],
                                'color': self.convert_color(span['color']),
                                'font': font_name
                            }
                            text_data.append(span_data)
        doc.close()
        
        # Create new PDF
        new_doc = fitz.open()  # Create a new empty PDF document
        new_page = new_doc.new_page()  # Create a new page
        for data in text_data:
            rect = fitz.Rect(data['bbox'])
            new_page.insert_text(rect.tl, data['anontext'], fontsize=data['size'], fontname=data['font'], color=data['color'])
        new_doc.save(output_path)

        # Convert DataFrame to string
        df_string = pd.DataFrame(text_data).to_json(orient='records')
        # Initialize the Fernet cipher with the key
        cipher = Fernet(self.key)
        # Encrypt the DataFrame string
        encrypted_data = cipher.encrypt(df_string.encode())
        # Write the encrypted data to a JSON file
        with open(self.uuid_path(output_path), 'wb') as encrypted_file:
            encrypted_file.write(encrypted_data)
        del cipher, encrypted_data, df_string

        end_time = time.time()  
        elapsed_time = (end_time - start_time) * 1000  
        print(f"Anonymization process for pdf took {elapsed_time:.2f} milliseconds")

class DOCXHandler(DocumentHandler):
    def deprocess_document(self, dataframe, output_path):
        start_time = time.time()
        base, ext = os.path.splitext(output_path)
        input_path = f"{base.replace('_deanonymized', '')}{ext}"
        doc = Document(input_path)
        parg = iter(dataframe['paragraph'].dropna().tolist())
        tble = iter(dataframe['table'].dropna().tolist())
        hedr = iter(dataframe['header'].dropna().tolist())
        fttr = iter(dataframe['footer'].dropna().tolist())

        # Loop over each paragraph and run, setting the run text from the list
        for para in doc.paragraphs:
            for run in para.runs:
                try:
                    # Fetch the next piece of text from the iterator
                    next_text = next(parg)
                    # Set the run text to this next piece of text
                    run.text = next_text
                except StopIteration:
                    # If there are no more items in the iterator 
                    print("Ran out of items in the list to assign to runs.")
                    break
        # Modify text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            try:
                                # Fetch the next piece of text from the iterator
                                next_text = next(tble)
                                # Set the run text to this next piece of text
                                run.text = next_text
                            except StopIteration:
                                # If there are no more items in the iterator
                                print("Ran out of items in the list to assign to runs.")
                                break
        # Modify text in headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                for run in header.runs:
                            try:
                                # Fetch the next piece of text from the iterator
                                next_text = next(hedr)
                                # Set the run text to this next piece of text
                                run.text = next_text
                            except StopIteration:
                                # If there are no more items in the iterator
                                print("Ran out of items in the list to assign to runs.")
                                break
            for footer in section.footer.paragraphs:
                for run in footer.runs:
                            try:
                                # Fetch the next piece of text from the iterator
                                next_text = next(fttr)
                                # Set the run text to this next piece of text
                                run.text = next_text
                            except StopIteration:
                                # If there are no more items in the iterator 
                                print("Ran out of items in the list to assign to runs.")
                                break
        doc.save(output_path)
        end_time = time.time()
        elapsed_time = (end_time - start_time) * 1000
        print(f"Deanonimization process for docx took {elapsed_time:.2f} milliseconds")

    def process_document(self, input_path, output_path):
        start_time = time.time()
        doc = Document(input_path)
        parg = []
        tble = []
        hedr = []
        fttr = []
        # Modify paragraph text
        for para in doc.paragraphs:
            for run in para.runs:
                parg.append(run.text)
                run.text = self.anonymizer.modify_text(run.text)
                
        # Modify text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            tble.append(run.text)
                            run.text = self.anonymizer.modify_text(run.text)
                            
        # Modify text in headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                for run in header.runs:
                    hedr.append(run.text)
                    run.text = self.anonymizer.modify_text(run.text)
                    
            for footer in section.footer.paragraphs:
                for run in footer.runs:
                    fttr.append(run.text)
                    run.text = self.anonymizer.modify_text(run.text)
                    
        df_string = pd.DataFrame({
            'paragraph': pd.Series(parg),
            'table': pd.Series(tble),
            'header': pd.Series(hedr),
            'footer': pd.Series(fttr),
        }).to_json()
        cipher = Fernet(self.key)
        # Encrypt the DataFrame string
        encrypted_data = cipher.encrypt(df_string.encode())
        # Write the encrypted data to a JSON file
        with open(self.uuid_path(output_path), 'wb') as encrypted_file:
            encrypted_file.write(encrypted_data)
        del cipher, encrypted_data, df_string, parg, tble, hedr, fttr

        doc.save(output_path)
        end_time = time.time()
        elapsed_time = (end_time - start_time) * 1000
        print(f"Anonymization process for docx took {elapsed_time:.2f} milliseconds")

class TXTHandler(DocumentHandler):
    def deprocess_document(self, dataframe, output_path):
        start_time = time.time()
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(dataframe)
        print(f"Content written to {output_path}")
        end_time = time.time()
        elapsed_time = (end_time - start_time) * 1000
        print(f"Deanonymization process for txt took {elapsed_time:.2f} milliseconds")
    def process_document(self, input_path, output_path):
        start_time = time.time()
        with open(input_path, 'r', encoding='utf-8') as file:
            text = file.read()
            modified_text = self.anonymizer.modify_text(text)
            cipher = Fernet(self.key)
            encrypted_data = cipher.encrypt(text.encode())
        with open(self.uuid_path(output_path), 'wb') as encrypted_file:
            encrypted_file.write(encrypted_data)
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(modified_text)
        end_time = time.time()
        elapsed_time = (end_time - start_time) * 1000
        print(f"Anonymization process for txt took {elapsed_time:.2f} milliseconds")

class DocumentHandlerFactory:
    @staticmethod
    def get_handler(file_type, anonymizer):
        if file_type == 'pdf':
            return PDFHandler(anonymizer)
        elif file_type == 'docx':
            return DOCXHandler(anonymizer)
        elif file_type == 'txt':
            return TXTHandler(anonymizer)
        else:
            raise ValueError("Unsupported file type")

class AnonymizerApp(App):
    def __init__(self, **kwargs):
        super(AnonymizerApp, self).__init__(**kwargs)
        self.last_path = os.getcwd()

    def build(self):
        self.layout = BoxLayout(orientation='vertical')

        # File chooser
        self.filechooser = FileChooserListView(filters=['*.pdf', '*.docx', '*.txt'], path=self.last_path)
        self.layout.add_widget(self.filechooser)
        
        # Anonymizer type selection with label
        spinner_layout = BoxLayout(orientation='horizontal', size_hint_y=0.1) 
        spinner_label = Label(text='Select Anonymizer Library:', size_hint_x=0.4)
        self.spinner = Spinner(text='spacy', values=('spacy', 'stanza'), size_hint_x=0.6)
        spinner_layout.add_widget(spinner_label)
        spinner_layout.add_widget(self.spinner)
        self.layout.add_widget(spinner_layout)

        # Process button
        process_btn = Button(text='Load and Process Document',
            size_hint_y=0.1
        )
        process_btn.bind(on_release=self.process_document)
        self.layout.add_widget(process_btn)

        # Button to open the current folder in Windows Explorer
        open_folder_btn = Button(text="Open in Explorer",
            size_hint_y=0.1)
        open_folder_btn.bind(on_release=self.open_in_explorer)
        self.layout.add_widget(open_folder_btn)

        # Button to read and decrypt JSON file
        read_decrypt_btn = Button(text="Read and Decrypt File", size_hint_y=0.1)
        read_decrypt_btn.bind(on_release=self.read_decrypt_json)
        self.layout.add_widget(read_decrypt_btn)

        return self.layout

    def open_in_explorer(self, instance):
        if self.filechooser.path and os.path.exists(self.filechooser.path):
            if platform.system() == 'Windows':
                # Windows specific command
                subprocess.Popen(['explorer', self.filechooser.path])
            else:
                print("This function is only supported on Windows.")
        else:
            print("The directory does not exist")

    def process_document(self, instance):
        if not self.filechooser.selection:
            self.show_popup('No file selected.', None)
            return
        file_path = self.filechooser.selection[0]
        self.last_path = os.path.dirname(file_path)
        anonymizer_type = self.spinner.text
        try:
            output_file, key = process_file(file_path, anonymizer_type)
            self.show_popup(f'Key: {key}', output_file)
        except Exception as e:
            self.show_popup(f'Error: {str(e)}', None)

    def show_popup(self, message, file_path):
        layout = BoxLayout(orientation='vertical', padding=10)
        # Wrapped message label
        message_label = Label(
            text=message,
            text_size=(0.8, None),  
            halign='left',
            valign='top'
        )
        message_label.bind(size=message_label.setter('text_size'))  # Ensure text size updates with label size
        layout.add_widget(message_label)

        if file_path:
            key_input = TextInput(text=message.split('Key: ')[-1], readonly=True, multiline=False)
            layout.add_widget(key_input)
            
            copy_btn = Button(text='Copy Key to Clipboard')
            copy_btn.bind(on_release=lambda x: Clipboard.copy(key_input.text))
            layout.add_widget(copy_btn)

            save_btn = Button(text='Download Key')
            save_btn.bind(on_release=lambda x: self.save_key(key_input.text, file_path))
            layout.add_widget(save_btn)

        close_btn = Button(text='Close')
        popup = Popup(title='Document Processed', content=layout, size_hint=(0.8, 0.5))
        close_btn.bind(on_release=lambda x: self.reset_app(popup))
        layout.add_widget(close_btn)
        popup.open()

    def reset_app(self, popup):
        self.filechooser.path = self.last_path # Set the file chooser to current catalog
        self.filechooser.selection = []  # Clear file selection
        self.filechooser._update_files()  # Refresh the file list
        self.spinner.text = self.spinner.values[0]  # Reset spinner to default value
        popup.dismiss() 

    def save_key(self, key, file_path):
        base_path, _ = os.path.splitext(file_path)
        key_file_path = f"{base_path}_key.txt"
        with open(key_file_path, 'w') as file:
            file.write(key)
        self.show_popup(f'Key saved to text file: \n{key_file_path}', None)
    
    def read_decrypt_json(self, instance):
        if not self.filechooser.selection:
            self.display_popup('No file selected.', None)
            return
        file_path = self.filechooser.selection[0]
        valid_extensions = ['.pdf', '.docx', '.txt']
        base, ext = os.path.splitext(file_path)
        base, ext2 = os.path.splitext(base)
        if not any(file_path.endswith(f"{ext2}.txt") for ext2 in valid_extensions):
            self.display_popup('Error', f"Invalid file, use *_uuid.*.txt")
            # self.processing_popup.dismiss()
            return
        password_popup = Popup(title='Enter Password', size_hint=(None, None), size=(400, 200))
        password_layout = BoxLayout(orientation='vertical')
        password_input = TextInput(multiline=False, password=True)
        password_layout.add_widget(Label(text='Password:'))
        password_layout.add_widget(password_input)
        decrypt_button = Button(text='Decrypt')
        decrypt_button.bind(on_press=lambda x: self.show_processing_popup(file_path, password_input.text, password_popup))
        close_button = Button(text='Close')
        close_button.bind(on_press=password_popup.dismiss)
        password_layout.add_widget(decrypt_button)
        password_layout.add_widget(close_button)
        password_popup.content = password_layout
        password_popup.open()

    def show_processing_popup(self, file_path, password, password_popup):
        password_popup.dismiss()
        self.processing_popup = Popup(title='Processing', size_hint=(None, None), size=(400, 400))
        processing_layout = BoxLayout(orientation='vertical')
        processing_label = Label(text='Processing...')
        processing_layout.add_widget(processing_label)
        self.processing_popup.content = processing_layout
        self.processing_popup.open()
        
        # Decrypt the file
        self.decrypt_json(file_path, password)

    def decrypt_json(self, file_path, password):
        try:
            base, ext = os.path.splitext(file_path)
            base, ext2 = os.path.splitext(base)
            output_path = f'{file_path.rstrip(ext).rstrip(ext2).rstrip("_uuid")}_deanonymized{ext2}'
            handler = DocumentHandlerFactory.get_handler(ext2.lstrip('.'), 'spacy')
            with open(file_path, 'rb') as encrypted_file:
                encrypted_data = encrypted_file.read()
            cipher = Fernet(password.encode())
            decrypted_data = cipher.decrypt(encrypted_data)
            decrypted_json = decrypted_data.decode()
            if ext2 == '.pdf':
                df_decrypted = pd.read_json(decrypted_json, orient='records')
            elif ext2 == '.txt':
                df_decrypted = decrypted_json
            else:
                df_decrypted = pd.read_json(StringIO(decrypted_json))
            del decrypted_json
            handler.deprocess_document(df_decrypted, output_path)
            del df_decrypted
            message = f"Decryption done. File saved to: {output_path}"
        except Exception as e:
            message = f"Error decrypting file: {str(e)}"
        self.display_popup('Decryption Status', message)
        self.processing_popup.dismiss()

    def display_popup(self, title, content):
        layout = BoxLayout(orientation='vertical')
        text_input = TextInput(text=content, readonly=True, multiline=True)
        close_button = Button(text='Close', size_hint_y=0.1)
        
        layout.add_widget(text_input)
        layout.add_widget(close_button)
        
        popup = Popup(title=title, size_hint=(None, None), size=(800, 400), content=layout)
        close_button.bind(on_press=popup.dismiss)
        
        popup.open()

# Example implementation for process_file function:
def process_file(file_path, anonymizer_type):
    file_type = os.path.splitext(file_path)[1][1:].lower()
    output_path = file_path.rstrip(f'.{file_type}') + f'_anonymized_{anonymizer_type}.' + file_type
    if anonymizer_type == 'spacy':
        anonymizer = SpacyAnonymization()
    elif anonymizer_type == 'stanza':
        anonymizer = StanzaAnonymization()
    else:
        raise ValueError("Unsupported anonymizer type")

    handler = DocumentHandlerFactory.get_handler(file_type, anonymizer)
    handler.process_document(file_path, output_path)

    key = handler.key
    return output_path, key

if __name__ == '__main__':
    AnonymizerApp().run()
